"""In-memory PDF/DOCX parser — no disk I/O, BytesIO only."""

import io
import logging
import re

import pdfplumber
from docx import Document
from fastapi import UploadFile

logger = logging.getLogger(__name__)

ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


def _normalize_whitespace(text: str) -> str:
    """Normalize whitespace while preserving Vietnamese diacritics."""
    # Collapse multiple spaces/tabs but keep newlines for structure
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse 3+ newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_pdf(data: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber (in-memory)."""
    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(layout=True) or ""
            if page_text.strip():
                pages.append(page_text)
    return _normalize_whitespace("\n\n".join(pages))


def parse_docx(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx (in-memory)."""
    doc = Document(io.BytesIO(data))
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table cells
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return _normalize_whitespace("\n\n".join(parts))


async def parse_file(file: UploadFile) -> str:
    """Validate and parse uploaded PDF or DOCX file. Returns extracted text."""
    filename = (file.filename or "").lower()
    content_type = file.content_type or ""

    # Validate extension
    ext = ""
    if filename.endswith(".pdf"):
        ext = ".pdf"
    elif filename.endswith(".docx"):
        ext = ".docx"

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type. Only PDF and DOCX are accepted.")

    # Read bytes into memory
    data = await file.read()

    # Validate size
    if len(data) > MAX_FILE_SIZE:
        raise ValueError(f"File too large. Maximum size is 10MB.")

    logger.info("Parsing %s file: %s (%d bytes)", ext, file.filename, len(data))

    if ext == ".pdf":
        return parse_pdf(data)
    else:
        return parse_docx(data)
